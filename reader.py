from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader


def read_pdf(file_path: str | Path) -> str:
    reader = PdfReader(str(file_path))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file_path: str | Path) -> str:
    document = Document(str(file_path))
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path: str | Path) -> str:
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return read_pdf(path)
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")


def load_job_description(file_path: str | Path | None = None) -> str:
    default_path = Path(__file__).resolve().parent / "job_description.txt"
    target_path = Path(file_path) if file_path is not None else default_path

    if not target_path.exists():
        raise FileNotFoundError(f"Job description file not found: {target_path}")

    if target_path.suffix.lower() == ".pdf":
        return read_pdf(target_path)
    if target_path.suffix.lower() == ".docx":
        return read_docx(target_path)
    return target_path.read_text(encoding="utf-8")
