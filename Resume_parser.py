import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)

model="openai/gpt-oss-120b" 
role="user"

# PART - 1:-
# ULOADING THE JOD DESCRIPTION GIVEN BY THE HR TO THE LLM AND EXTRACTING THE RELEVANT INFORMATION FROM IT.


#Job description provided by the HR for the role of Junior Instrumentation & Control Engineer. The job description includes the job title, department, reporting structure, job summary, key responsibilities, and required qualifications.

job_description="""
Description:
Job Title: Junior Instrumentation & Control EngineerDepartment: Engineering / Operations & MaintenanceReports to: Senior Instrumentation Engineer / Project Manager
Job Summary:
The Junior Instrumentation & Control Engineer will assist in the design, installation, calibration, and maintenance of measurement and control systems. The successful candidate will work under the guidance of senior engineers to ensure that plant processes operate safely, efficiently, and in compliance with industry regulations. 
Key Responsibilities:
System Design Support: Assist in the preparation of technical documentation, including P&IDs (Piping and Instrumentation Diagrams), instrument lists, I/O lists, and wiring diagrams.  
Installation & Commissioning: Support the on-site installation of field instruments (sensors, transmitters, control valves) and participate in the commissioning/start-up of new systems.  
Calibration & Testing: Perform routine calibration of pressure, temperature, flow, and level instruments. Assist in loop checks and functional testing of control loops. 
Maintenance & Troubleshooting: Help troubleshoot hardware and software issues within PLC, DCS, and SCADA systems. Identify root causes for equipment malfunctions.  
Data Analysis: Use CMMS (Computerized Maintenance Management Systems) or spreadsheets to log maintenance activities, track instrument failure rates, and update system parameters.  
Safety & Compliance: Ensure all work is performed in accordance with safety standards (e.g., ISO, IEC, OSHA) and internal company policies.  
Required Qualifications
Education: Bachelor’s degree in Instrumentation Engineering, Electronics and Communication Engineering, or a related field.  

Technical Knowledge: Foundational understanding of control theory, PID loops, and field instrumentation.  
Tools/Software: Basic familiarity with CAD software (e.g., AutoCAD), Microsoft Office (Excel), and diagnostic tools (multimeters, oscilloscopes).  
Soft Skills: Strong analytical mindset, attention to detail, and effective communication skills for reporting technical issues.  
""" 

#Creating a Pydantic model for the job description. The model will have fields for the role, required skills, preferred skills, minimum experience, education requirements and responsibilities.
class JobDescription(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

#Creating a JSON schema for the JobDescription model using Pydantic's model_json_schema() method. This schema will be used to validate the extracted information from the job description.
jobd_schema = JobDescription.model_json_schema()


#System prompt for the LLM to extract the relevant information from the job description and return it in a structured JSON format that adheres to the JobDescription schema.
system_prompt = f"""
You are a highly skilled and experienced recruiter specializing in the field of Instrumentation & Control Engineering.
Return the job description in a structured JSON format that adheres to the following schema:
{jobd_schema}

IMPORTANT:
Do NOT return the schema itself.
Don NOT return field like "properties", "type", "title", "description", "required" or any other schema-related metadata.
Fill the schema  with the relevant information extracted from the provided job description.

If minimum_experience is not mentioned in the job description, return it as null.
If infromation is not available for any field, return it as an empty list or null as appropriate.
Do not invent any information that is not present in the job description.
"""

#User prompt given by the recruiter for the LLM to analyze the job description and extract the relevant information to populate the fields of the JobDescription schema.    
user_prompt = f""" 
Analyse the following job description and extract the relevant information to populate the fields of the schema.

{job_description}
"""
#Messages to be sent to the LLM for processing. The system message contains the system prompt, and the user message contains the user prompt. The response format is specified as a JSON object.

message_system = {
    "role": "system", 
    "content": system_prompt
}
#
message_user = {
    "role": "user",
    "content": user_prompt
}
response_format = {
    "type": "json_object",
}

#This is the part which will send the details to the LLM and give us the answer of the response and we will convert that answer into raw_json.
Messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=Messages, response_format=response_format)

answer=response.choices[0].message.content

raw_json = answer
#print(raw_json)


import json              #Importing the json module to parse the JSON response from the LLM and create a JobDescription object from it.
job_data = json.loads(raw_json)

job = JobDescription(**job_data)

# PART - 2:-
# ULOADING THE RESUME OF THE CANDIDATE TO THE LLM AND EXTRACTING THE RELEVANT INFORMATION FROM IT.


#This is for the final detailing of the resume, here we will get the final score ad the details for that score (means why the candidate has been given taht score)
class MatchResult(BaseModel):   #section of part 4
    score : float
    details : dict

#Creating a Pydantic model for the experience of the candidate. The model will have fields for the company name, job title, start date, role, duration, responsibilities, designation, description and skills used.
class Experience(BaseModel):
    company_name: str | None = None
    job_title: str | None = None
    start_date: str | None = None
    role: str | None = None
    duration: str | None = None
    responsibilities: list[str] | None = None
    designation: str | None = None
    description: str | None = None
    skills_used: list[str] | None = None


#Creating a Pydantic model for the resume of the candidate. The model will have fields for the name, email, phone number, linkedin profile, career objective, education, skills, technical skills, projects, certifications, achievements, total experience years and experience. The experience field will be a list of Experience objects.
class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone_number: str | None = None
    linkedin_profile: str | None = None
    career_objective: str | None = None
    education: list[str] | None = None
    skills: list[str] | None = None
    technical_skills: list[str] | None = None
    projects: list[str] | None = None
    certifications: list[str] | None = None
    achievements: list[str] | None = None
    total_experience_years: str | None = None
    experience: list[Experience] | None = None

# create JSON schema for Resume model
resume_schema = Resume.model_json_schema()

#This is also the section of PART 4 , and its tastk is to give the HR the final versict about the Resume and Candidate (in short why they have given the point, there what skills meet with the requirements of the company, what is missing from there resume and there ther overall percentage %. )
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESSCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}

    Return JSON matching this schema:
    {match_schema}
 
    Give me :
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Wheather experince requirement is met
    5. Overall match percentage from 0 to 100
    6. A short finall verdict

    Keep theresponse concise and easy to read.
"""
    messages = [{
        "role": "user",
        "content": prompt
    }]
    response_format = {
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    return MatchResult(**data)

#Here in this part i am converting the normal resume text into the JSON format
def parse_resume(resume_text):      #section of part 4
    system_prompt = f"""
You are an expert resume parser.

Exxtract information from the resume based on its meaning,
not only based on exact section headings.

Different resumes may use different headings.

For example:
 - Experince
 - Professional Experience
 - Work History
 - Employment
 - Internships

These maya ll contain relevant experince.

Skills may also appear in the skills section, work, experince,
internships or projects.

Return ONLY valid JSON matching this schema:

{resume_schema}
Important Rules:

1. Do not invent information.
2. If a value is not available, return null.
3. If a list has no information, return an empty list.
4. Include internships inside experinces.
5. Extract skills mentioned across the entire resume.
"""
    
    user_prompt = f"""
Parse the following resume:

{resume_text}
"""
    message_system = {
        "role" : "system",
        "content" : system_prompt
    }
    message_user = {
        "role" : "user",
        "content" : user_prompt
    }
    messages = [message_system, message_user]
    response_format = {
        "type" : "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

# PART - 3:-
# READING THE RESUME OF THE CANDIDATE FROM A FILE AND EXTRACTING THE RELEVANT INFORMATION FROM IT.
# READING THE RESUME OP;OAED IN ONLY 2 FORMATS - PDF AND DOCX. IF THE RESUME IS IN ANY OTHER FORMAT, IT WILL NOT BE READ.


#PDFREADER AND DOCX READER FUNCTIONS
#pdfReader and functions to read the resume files in PDF formats respectively.
from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text() 
        if page_text:
            text += page_text + "\n"
    return text


#docx reader function to read the resume files in DOCX format and extract the text from them.
def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"


#if the resume is in docx format, we will also read the tables in the docx file and extract the text from them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


#Time to read the resume file and extract the text from it. We will check the file extension and call the appropriate function to read the file.
def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")
    


#PART - 4:-
# Now we will parse the actual resume of the candidate and extract the relevant information from it using the LLM. We will use the same approach as we did for the job description.   


resume_folder = Path("RESUMES")               #For finding the resumes from the folder.
all_results=[]                                # For getting all the results because we have more than 1 resume. 
for file_path in resume_folder.iterdir():      #This will add the file path of the file into the systems.
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
       continue 
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)
    time.sleep(5)
    result = final_score(job, parsed_resume)
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name" : parsed_resume.name,
        "score" : result.score,
        "details" : result.details
    })

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)

top_candidate  = all_results[:1]
worst_candidate = all_results[-1:]

print("TOP CANDIDATE")
for candidate in top_candidate:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])

print("LOWEST CANDIDATE")
for candidate  in worst_candidate:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])    